
import streamlit as st
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from collections import defaultdict
from io import BytesIO
import xlsxwriter
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import networkx as nx

st.set_page_config(page_title="ProjectPERT", page_icon="◈", layout="wide", initial_sidebar_state="expanded")

# Force sidebar to always stay open
st.markdown("""
<style>
section[data-testid="stSidebar"] { min-width: 300px !important; width: 300px !important; transform: none !important; }
button[data-testid="collapsedControl"] { display: none !important; }
[data-testid="stSidebarCollapseButton"] { display: none !important; }
</style>
""", unsafe_allow_html=True)

st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Playfair+Display:wght@600;700&family=Inter:wght@300;400;500;600&family=JetBrains+Mono:wght@400;600&display=swap');
html, body, [class*="css"] { font-family: 'Inter', sans-serif; font-size: 15px; color: #1a2332; }
.stApp { background: #f4f9fb; }
#MainMenu, footer, header { visibility: hidden; }
.hero { background: linear-gradient(135deg, #0a7e8c 0%, #0d9eaf 50%, #14b8cc 100%); border-radius: 20px; padding: 3rem 3.5rem; margin-bottom: 2rem; position: relative; overflow: hidden; box-shadow: 0 8px 32px rgba(10,126,140,0.25); }
.hero-eyebrow { font-family: 'JetBrains Mono', monospace; font-size: 0.72rem; letter-spacing: 4px; text-transform: uppercase; color: rgba(255,255,255,0.65); margin-bottom: 0.8rem; }
.hero-title { font-family: 'Playfair Display', serif; font-size: 3rem; font-weight: 700; color: #ffffff; margin: 0 0 0.6rem 0; line-height: 1.15; }
.hero-sub { color: rgba(255,255,255,0.8); font-size: 1.05rem; font-weight: 300; margin: 0; max-width: 600px; line-height: 1.6; }
.hero-badges { display: flex; gap: 0.6rem; margin-top: 1.5rem; flex-wrap: wrap; }
.hero-badge { background: rgba(255,255,255,0.15); border: 1px solid rgba(255,255,255,0.25); color: white; border-radius: 20px; padding: 0.3rem 0.9rem; font-size: 0.78rem; font-weight: 500; }
.card { background: #ffffff; border: 1px solid #e0eef1; border-radius: 16px; padding: 1.6rem 1.8rem; margin-bottom: 1.2rem; box-shadow: 0 2px 12px rgba(10,126,140,0.06); }
.card-accent { border-left: 4px solid #0a7e8c; }
.section-label { font-family: 'JetBrains Mono', monospace; font-size: 0.68rem; letter-spacing: 3.5px; text-transform: uppercase; color: #0a7e8c; margin-bottom: 0.8rem; }
.metric-grid { display: grid; grid-template-columns: repeat(4, 1fr); gap: 1rem; margin-bottom: 1.5rem; }
.metric-tile { background: #ffffff; border: 1px solid #e0eef1; border-radius: 14px; padding: 1.4rem 1.2rem; text-align: center; box-shadow: 0 2px 10px rgba(10,126,140,0.06); position: relative; overflow: hidden; }
.metric-tile::before { content: ''; position: absolute; top: 0; left: 0; right: 0; height: 3px; background: linear-gradient(90deg, #0a7e8c, #14b8cc); }
.metric-val { font-family: 'Playfair Display', serif; font-size: 2.4rem; font-weight: 700; color: #0a7e8c; line-height: 1.1; }
.metric-unit { font-size: 0.85rem; color: #6b8a92; font-weight: 400; }
.metric-lbl { font-size: 0.8rem; color: #7a9aa4; margin-top: 0.4rem; font-weight: 500; }
.critical-path-box { background: linear-gradient(135deg, #fff5f5 0%, #fff8f8 100%); border: 1px solid #fecaca; border-left: 4px solid #ef4444; border-radius: 12px; padding: 1.2rem 1.5rem; margin-bottom: 1.2rem; }
.critical-path-label { font-family: 'JetBrains Mono', monospace; font-size: 0.68rem; letter-spacing: 3px; text-transform: uppercase; color: #ef4444; margin-bottom: 0.6rem; }
.critical-path-text { font-size: 1rem; color: #b91c1c; font-weight: 600; line-height: 1.8; }
.info-box { background: linear-gradient(135deg, #f0fafb 0%, #e8f6f8 100%); border: 1px solid #b2dde3; border-left: 4px solid #0a7e8c; border-radius: 12px; padding: 1.2rem 1.5rem; margin-bottom: 1rem; font-size: 0.92rem; color: #1a4a52; line-height: 1.7; }
section[data-testid="stSidebar"] { background: #ffffff !important; border-right: 1px solid #e0eef1; }
section[data-testid="stSidebar"] label { color: #2a4a52 !important; font-size: 0.85rem !important; font-weight: 600 !important; }
.stButton>button { background: linear-gradient(135deg, #0a7e8c, #0d9eaf); color: white; border: none; border-radius: 10px; padding: 0.7rem 1.8rem; font-size: 0.92rem; font-weight: 600; box-shadow: 0 3px 10px rgba(10,126,140,0.25); }
.stButton>button:hover { opacity: 0.88; }
.stDownloadButton>button { background: linear-gradient(135deg, #0a7e8c, #0d9eaf) !important; color: white !important; border: none !important; border-radius: 10px !important; font-weight: 600 !important; }
div[data-testid="stNumberInput"] input, div[data-testid="stTextInput"] input { background: #f8fcfd !important; border: 1.5px solid #c8e6eb !important; border-radius: 8px !important; font-size: 0.92rem !important; color: #1a2332 !important; }
.stTabs [data-baseweb="tab-list"] { gap: 0.5rem; background: transparent; border-bottom: 2px solid #e0eef1; }
.stTabs [data-baseweb="tab"] { font-weight: 600; font-size: 0.88rem; color: #6b8a92; border-radius: 8px 8px 0 0; padding: 0.6rem 1.2rem; }
.stTabs [aria-selected="true"] { color: #0a7e8c !important; border-bottom: 3px solid #0a7e8c; }
div[data-testid="stDataFrame"] { border-radius: 12px; overflow: hidden; border: 1px solid #e0eef1; }
hr { border-color: #e0eef1; margin: 1.5rem 0; }
p { font-size: 0.92rem; line-height: 1.7; }
.guide-step { display: flex; gap: 1rem; align-items: flex-start; margin-bottom: 1rem; }
.guide-num { background: linear-gradient(135deg, #0a7e8c, #14b8cc); color: white; width: 28px; height: 28px; border-radius: 50%; display: flex; align-items: center; justify-content: center; font-size: 0.78rem; font-weight: 700; flex-shrink: 0; margin-top: 2px; }
.guide-text { font-size: 0.9rem; color: #2a4a52; line-height: 1.6; }
</style>
""", unsafe_allow_html=True)

UNIT_OPTIONS = ["Minutes", "Hours", "Days", "Weeks", "Months", "Years"]
TEAL = "#0a7e8c"
RED = "#ef4444"

def pert_mean(a, m, b): return (a + 4*m + b) / 6
def pert_std(a, b): return (b - a) / 6

def topological_sort(activities):
    in_degree = {act: 0 for act in activities}
    adj = defaultdict(list)
    for act, info in activities.items():
        for pred in info["predecessors"]:
            adj[pred].append(act)
            in_degree[act] += 1
    queue = [a for a in activities if in_degree[a] == 0]
    order = []
    while queue:
        node = queue.pop(0)
        order.append(node)
        for neighbor in adj[node]:
            in_degree[neighbor] -= 1
            if in_degree[neighbor] == 0:
                queue.append(neighbor)
    return order

def simulate_project(activities, n_simulations=10_000):
    order = topological_sort(activities)
    project_durations = []
    for _ in range(n_simulations):
        finish = {}
        for act in order:
            a, m, b = activities[act]["min"], activities[act]["avg"], activities[act]["max"]
            duration = np.random.triangular(a, m, b)
            preds = activities[act]["predecessors"]
            start = max((finish[p] for p in preds), default=0)
            finish[act] = start + duration
        project_durations.append(max(finish.values()))
    return np.array(project_durations)

def compute_pert_critical_path(activities):
    order = topological_sort(activities)
    exp_dur = {a: pert_mean(activities[a]["min"], activities[a]["avg"], activities[a]["max"]) for a in activities}
    earliest_finish = {}
    for act in order:
        preds = activities[act]["predecessors"]
        start = max((earliest_finish[p] for p in preds), default=0)
        earliest_finish[act] = start + exp_dur[act]
    end_node = max(earliest_finish, key=earliest_finish.get)
    def backtrack(node):
        preds = activities[node]["predecessors"]
        if not preds: return [node]
        return backtrack(max(preds, key=lambda p: earliest_finish[p])) + [node]
    return max(earliest_finish.values()), backtrack(end_node), exp_dur, earliest_finish

def draw_network_diagram(activities, critical_path):
    G = nx.DiGraph()
    for act in activities: G.add_node(act)
    for act, info in activities.items():
        for pred in info["predecessors"]: G.add_edge(pred, act)
    order = topological_sort(activities)
    levels = {}
    for node in order:
        preds = activities[node]["predecessors"]
        levels[node] = max((levels[p] for p in preds), default=-1) + 1
    level_counts = defaultdict(int)
    level_y = defaultdict(int)
    for node in order: level_counts[levels[node]] += 1
    pos = {}
    for node in order:
        lv = levels[node]; count = level_counts[lv]; idx = level_y[lv]
        pos[node] = (lv * 2.5, (count - 1) / 2.0 - idx * 1.6)
        level_y[lv] += 1
    fig, ax = plt.subplots(figsize=(12, max(4, len(activities) * 0.85)))
    fig.patch.set_facecolor("#f4f9fb"); ax.set_facecolor("#f4f9fb")
    crit_set = set(critical_path)
    crit_edges = [(critical_path[i], critical_path[i+1]) for i in range(len(critical_path)-1)]
    non_crit_edges = [(u, v) for u, v in G.edges() if (u, v) not in crit_edges]
    nx.draw_networkx_edges(G, pos, edgelist=non_crit_edges, ax=ax, edge_color="#b2c8cc", width=1.8, arrows=True, arrowsize=20, connectionstyle="arc3,rad=0.05", arrowstyle="-|>", min_source_margin=22, min_target_margin=22)
    nx.draw_networkx_edges(G, pos, edgelist=crit_edges, ax=ax, edge_color=RED, width=3.0, arrows=True, arrowsize=25, connectionstyle="arc3,rad=0.05", arrowstyle="-|>", min_source_margin=22, min_target_margin=22)
    nx.draw_networkx_nodes(G, pos, nodelist=[n for n in G.nodes() if n not in crit_set], ax=ax, node_color="#e8f5f7", node_size=1800, edgecolors=TEAL, linewidths=2)
    nx.draw_networkx_nodes(G, pos, nodelist=list(crit_set), ax=ax, node_color=RED, node_size=2000, edgecolors="#b91c1c", linewidths=2)
    for node, (x, y) in pos.items():
        is_crit = node in crit_set
        ax.text(x, y + 0.12, node, ha="center", va="center", fontsize=13, fontweight="bold", color="white" if is_crit else TEAL)
        name = activities[node]["name"]; short = name if len(name) <= 14 else name[:12] + "..."
        ax.text(x, y - 0.28, short, ha="center", va="center", fontsize=7.5, color="#4a6a72")
    ax.legend(handles=[mpatches.Patch(color=RED, label="Critical Path"), mpatches.Patch(facecolor="#e8f5f7", edgecolor=TEAL, linewidth=2, label="Non-Critical")], loc="upper left", framealpha=0.9, fontsize=9, facecolor="white", edgecolor="#e0eef1")
    ax.set_title("Project Network Diagram", fontsize=14, fontweight="bold", color="#1a2332", pad=15)
    ax.axis("off"); fig.tight_layout(pad=1.5)
    return fig

def set_cell_bg(cell, hex_color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr(); shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hex_color); tcPr.append(shd)

def export_excel(activities, durations, pert_dur, crit_path, sl, unit):
    output = BytesIO()
    mean_dur = float(np.mean(durations)); std_dur = float(np.std(durations)); sl_val = float(np.percentile(durations, sl))
    wb = xlsxwriter.Workbook(output, {"in_memory": True})
    title_fmt = wb.add_format({"bold": True, "font_size": 16, "font_color": "#0a7e8c", "bottom": 2, "bottom_color": "#0a7e8c", "font_name": "Calibri"})
    hdr_fmt   = wb.add_format({"bold": True, "bg_color": "#0a7e8c", "font_color": "#ffffff", "font_size": 11, "border": 1, "align": "center", "font_name": "Calibri"})
    cell_fmt  = wb.add_format({"font_size": 11, "border": 1, "border_color": "#d0e8eb", "font_name": "Calibri"})
    num_fmt   = wb.add_format({"font_size": 11, "border": 1, "border_color": "#d0e8eb", "num_format": "0.00", "align": "center", "font_name": "Calibri"})
    hi_fmt    = wb.add_format({"font_size": 11, "border": 1, "bg_color": "#fff0f0", "font_color": "#c00000", "bold": True, "font_name": "Calibri"})
    ml_fmt    = wb.add_format({"bold": True, "font_size": 12, "bg_color": "#e8f6f8", "border": 1, "font_name": "Calibri"})
    mv_fmt    = wb.add_format({"font_size": 12, "border": 1, "num_format": "0.00", "font_color": "#0a7e8c", "bold": True, "font_name": "Calibri"})
    red_fmt   = wb.add_format({"font_size": 11, "border": 1, "font_color": "#c00000", "bold": True, "font_name": "Calibri"})
    ws1 = wb.add_worksheet("Summary"); ws1.set_column("A:A", 36); ws1.set_column("B:B", 24)
    ws1.write("A1", "ProjectPERT - Simulation Summary", title_fmt); ws1.write("A2", f"Time unit: {unit}")
    for i, (lbl, val) in enumerate([("Mean Project Duration", mean_dur), ("Standard Deviation", std_dur), (f"{sl}% Service Level Duration", sl_val), ("PERT Expected Duration", pert_dur)], start=3):
        ws1.write(i, 0, lbl, ml_fmt); ws1.write(i, 1, val, mv_fmt)
    ws1.write(7, 0, "Critical Path", ml_fmt); ws1.write(7, 1, " > ".join(crit_path), red_fmt)
    ws2 = wb.add_worksheet("Activities"); ws2.set_column("A:A", 8); ws2.set_column("B:B", 28); ws2.set_column("C:C", 22); ws2.set_column("D:I", 14)
    for col, h in enumerate(["Label","Activity","Predecessors",f"Min ({unit})",f"Avg ({unit})",f"Max ({unit})","PERT Mean","Std Dev","Critical?"]):
        ws2.write(0, col, h, hdr_fmt)
    for row, (lbl, info) in enumerate(activities.items(), start=1):
        mu = pert_mean(info["min"], info["avg"], info["max"]); sig = pert_std(info["min"], info["max"])
        fmt = hi_fmt if lbl in crit_path else cell_fmt
        ws2.write(row, 0, lbl, fmt); ws2.write(row, 1, info["name"], fmt); ws2.write(row, 2, ", ".join(info["predecessors"]) or "-", fmt)
        ws2.write(row, 3, info["min"], num_fmt); ws2.write(row, 4, info["avg"], num_fmt); ws2.write(row, 5, info["max"], num_fmt)
        ws2.write(row, 6, round(mu, 2), num_fmt); ws2.write(row, 7, round(sig, 2), num_fmt); ws2.write(row, 8, "YES" if lbl in crit_path else "", fmt)
    ws3 = wb.add_worksheet("Simulation Data"); ws3.set_column("A:B", 20)
    ws3.write(0, 0, "Simulation #", hdr_fmt); ws3.write(0, 1, f"Duration ({unit})", hdr_fmt)
    for i, val in enumerate(durations[:5000], start=1): ws3.write(i, 0, i, cell_fmt); ws3.write(i, 1, round(float(val), 4), num_fmt)
    ws4 = wb.add_worksheet("Histogram"); ws4.set_column("A:B", 22)
    ws4.write(0, 0, f"Duration Bin ({unit})", hdr_fmt); ws4.write(0, 1, "Frequency", hdr_fmt)
    counts, bin_edges = np.histogram(durations, bins=40)
    for i, (count, edge) in enumerate(zip(counts, bin_edges), start=1): ws4.write(i, 0, round(float(edge), 2), num_fmt); ws4.write(i, 1, int(count), cell_fmt)
    chart = wb.add_chart({"type": "column"})
    chart.add_series({"name": "Frequency", "categories": ["Histogram", 1, 0, len(counts), 0], "values": ["Histogram", 1, 1, len(counts), 1], "fill": {"color": "#0a7e8c"}, "border": {"color": "#ffffff"}})
    chart.set_title({"name": "Project Duration Distribution"}); chart.set_x_axis({"name": f"Duration ({unit})"}); chart.set_y_axis({"name": "Frequency"}); chart.set_style(10)
    ws4.insert_chart("D2", chart, {"x_scale": 1.8, "y_scale": 1.5})
    wb.close(); output.seek(0); return output

def export_word(activities, durations, pert_dur, crit_path, sl, unit):
    mean_dur = float(np.mean(durations)); std_dur = float(np.std(durations)); sl_val = float(np.percentile(durations, sl))
    doc = DocxDocument()
    for section in doc.sections:
        section.top_margin = Cm(2.2); section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.5); section.right_margin = Cm(2.5)
    title = doc.add_paragraph(); run = title.add_run("ProjectPERT")
    run.bold = True; run.font.size = Pt(28); run.font.color.rgb = RGBColor(0x0a, 0x7e, 0x8c); run.font.name = "Calibri"
    sub = doc.add_paragraph(); sub_run = sub.add_run("Simulation Report - Project Duration Forecasting")
    sub_run.font.size = Pt(12); sub_run.font.color.rgb = RGBColor(0x6b, 0x8a, 0x92); sub_run.font.name = "Calibri"
    doc.add_paragraph()
    def add_heading(text):
        p = doc.add_paragraph(); run = p.add_run(text); run.bold = True; run.font.name = "Calibri"
        run.font.size = Pt(14); run.font.color.rgb = RGBColor(0x0a, 0x7e, 0x8c)
        p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(4)
    def add_body(text):
        p = doc.add_paragraph(text)
        for r in p.runs: r.font.name = "Calibri"; r.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(4)
    add_heading("1. Executive Summary")
    add_body(f"This report presents Monte Carlo PERT simulation results. All durations are in {unit.lower()}.")
    doc.add_paragraph()
    tbl = doc.add_table(rows=5, cols=2); tbl.style = "Table Grid"
    for i, h in enumerate(["Metric", "Value"]):
        cell = tbl.rows[0].cells[i]; cell.text = h; run = cell.paragraphs[0].runs[0]
        run.bold = True; run.font.name = "Calibri"; run.font.size = Pt(11); run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        set_cell_bg(cell, "0a7e8c")
    for row_idx, (label, value) in enumerate([("Mean Project Duration", f"{mean_dur:.2f} {unit}"), ("Standard Deviation", f"{std_dur:.2f} {unit}"), (f"{sl}% Service Level", f"{sl_val:.2f} {unit}"), ("PERT Expected Duration", f"{pert_dur:.2f} {unit}")], start=1):
        row = tbl.rows[row_idx]; row.cells[0].text = label; row.cells[1].text = value
        for cell in row.cells:
            for r in cell.paragraphs[0].runs: r.font.name = "Calibri"; r.font.size = Pt(11)
        if row_idx % 2 == 0:
            for cell in row.cells: set_cell_bg(cell, "e8f6f8")
    doc.add_paragraph(); add_heading("2. Critical Path")
    p = doc.add_paragraph(); run = p.add_run(" > ".join(f"{l} ({activities[l]['name']})" for l in crit_path))
    run.bold = True; run.font.name = "Calibri"; run.font.size = Pt(11); run.font.color.rgb = RGBColor(0xC0,0x00,0x00)
    doc.add_paragraph(); add_heading("3. Activity Details")
    cols = ["Label","Activity","Predecessors",f"Min ({unit})",f"Avg ({unit})",f"Max ({unit})","PERT Mean","Std Dev","Critical?"]
    atbl = doc.add_table(rows=len(activities)+1, cols=len(cols)); atbl.style = "Table Grid"
    for col_idx, col_name in enumerate(cols):
        cell = atbl.rows[0].cells[col_idx]; cell.text = col_name; run = cell.paragraphs[0].runs[0]
        run.bold = True; run.font.name = "Calibri"; run.font.size = Pt(10); run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        set_cell_bg(cell, "0a7e8c")
    for row_idx, (lbl, info) in enumerate(activities.items(), start=1):
        mu = pert_mean(info["min"], info["avg"], info["max"]); sig = pert_std(info["min"], info["max"]); is_crit = lbl in crit_path
        vals = [lbl, info["name"], ", ".join(info["predecessors"]) or "-", str(info["min"]), str(info["avg"]), str(info["max"]), f"{mu:.2f}", f"{sig:.2f}", "YES" if is_crit else ""]
        for col_idx, val in enumerate(vals):
            cell = atbl.rows[row_idx].cells[col_idx]; cell.text = val; run = cell.paragraphs[0].runs[0]
            run.font.name = "Calibri"; run.font.size = Pt(10)
            if is_crit: run.bold = True; run.font.color.rgb = RGBColor(0xC0,0x00,0x00); set_cell_bg(cell, "fff0f0")
    doc.add_paragraph(); add_heading("4. Interpretation")
    for bullet in [f"Mean project duration: {mean_dur:.1f} {unit.lower()}.", f"50% chance of completing within {float(np.percentile(durations,50)):.1f} {unit.lower()}.", f"For {sl}% service level, plan for {sl_val:.1f} {unit.lower()}.", f"Std deviation of {std_dur:.1f} {unit.lower()} reflects schedule uncertainty."]:
        p = doc.add_paragraph(style="List Bullet"); run = p.add_run(bullet); run.font.name = "Calibri"; run.font.size = Pt(11)
    output = BytesIO(); doc.save(output); output.seek(0); return output

if "activities" not in st.session_state:
    st.session_state.activities = {
        "A": {"name": "Design",                "predecessors": [],        "min": 16, "avg": 21, "max": 26},
        "B": {"name": "Build Prototype",       "predecessors": ["A"],     "min": 3,  "avg": 6,  "max": 9},
        "C": {"name": "Evaluate Equipment",    "predecessors": ["A"],     "min": 5,  "avg": 7,  "max": 9},
        "D": {"name": "Test Prototype",        "predecessors": ["B"],     "min": 2,  "avg": 3,  "max": 4},
        "E": {"name": "Write Equipment Report","predecessors": ["C","D"], "min": 4,  "avg": 6,  "max": 8},
        "F": {"name": "Write Methods Report",  "predecessors": ["C","D"], "min": 6,  "avg": 8,  "max": 10},
        "G": {"name": "Write Final Report",    "predecessors": ["E","F"], "min": 1,  "avg": 2,  "max": 3},
    }
if "results"   not in st.session_state: st.session_state.results   = None
if "time_unit" not in st.session_state: st.session_state.time_unit = "Weeks"

with st.sidebar:
    st.markdown('<div style="padding:1rem 0 0.5rem;"><div style="font-size:0.65rem;letter-spacing:3px;text-transform:uppercase;color:#0a7e8c;margin-bottom:0.3rem;">Navigation</div><div style="font-size:1.2rem;font-weight:700;color:#1a2332;">ProjectPERT</div></div>', unsafe_allow_html=True)
    st.markdown("---")
    page = st.radio("", ["Dashboard", "User Guide", "About PERT"], label_visibility="collapsed")
    st.markdown("---")
    st.markdown('<div style="font-size:0.78rem;font-weight:600;color:#0a7e8c;letter-spacing:1px;text-transform:uppercase;margin-bottom:0.4rem;">Time Unit</div>', unsafe_allow_html=True)
    time_unit = st.selectbox("", UNIT_OPTIONS, index=UNIT_OPTIONS.index(st.session_state.time_unit), label_visibility="collapsed")
    st.session_state.time_unit = time_unit
    st.markdown("---")
    st.markdown('<div style="font-size:0.78rem;font-weight:600;color:#0a7e8c;letter-spacing:1px;text-transform:uppercase;margin-bottom:0.4rem;">Simulation Settings</div>', unsafe_allow_html=True)
    n_sim = st.select_slider("Iterations", options=[1_000, 5_000, 10_000, 50_000, 100_000], value=10_000, format_func=lambda x: f"{x:,}")
    service_level = st.slider("Service Level (%)", min_value=50, max_value=99, value=95, step=1)
    st.markdown("---")
    run_btn = st.button("Run Simulation", use_container_width=True)
    st.markdown('<div style="margin-top:1.5rem;padding:1rem;background:#f0fafb;border-radius:10px;border:1px solid #b2dde3;font-size:0.78rem;color:#2a6a74;line-height:1.6;"><strong>Tip</strong><br>Add activities on the left or upload a CSV/Excel, then click Run Simulation.</div>', unsafe_allow_html=True)

st.markdown('''<div class="hero">
  <div class="hero-eyebrow">Project Management Analytics</div>
  <div class="hero-title">ProjectPERT</div>
  <p class="hero-sub">Intelligent project duration forecasting using PERT analysis and Monte Carlo simulation. Understand uncertainty, identify critical paths, and plan with confidence.</p>
  <div class="hero-badges">
    <span class="hero-badge">Monte Carlo Simulation</span>
    <span class="hero-badge">Critical Path Analysis</span>
    <span class="hero-badge">PERT Estimation</span>
    <span class="hero-badge">Excel and Word Export</span>
  </div>
</div>''', unsafe_allow_html=True)

if page == "User Guide":
    st.markdown('<div class="section-label">User Guide</div>', unsafe_allow_html=True)
    col1, col2 = st.columns(2, gap="large")
    with col1:
        st.markdown('''<div class="card card-accent">
          <div class="section-label">Getting Started</div>
          <div class="guide-step"><div class="guide-num">1</div><div class="guide-text"><strong>Choose your time unit</strong> from the sidebar.</div></div>
          <div class="guide-step"><div class="guide-num">2</div><div class="guide-text"><strong>Enter activities</strong> manually or upload a CSV/Excel file.</div></div>
          <div class="guide-step"><div class="guide-num">3</div><div class="guide-text"><strong>Set predecessors</strong> — activities that must finish before the next can start.</div></div>
          <div class="guide-step"><div class="guide-num">4</div><div class="guide-text"><strong>Set service level</strong> and iterations in the sidebar.</div></div>
          <div class="guide-step"><div class="guide-num">5</div><div class="guide-text"><strong>Click Run Simulation</strong> to see results, histogram, and network diagram.</div></div>
          <div class="guide-step"><div class="guide-num">6</div><div class="guide-text"><strong>Export</strong> your results as Excel or Word.</div></div>
        </div>''', unsafe_allow_html=True)
    with col2:
        st.markdown('<div class="card card-accent"><div class="section-label">CSV/Excel Upload Format</div><p>Your file must have these columns:</p></div>', unsafe_allow_html=True)
        st.dataframe(pd.DataFrame({"Column": ["label","name","predecessors","min","avg","max"], "Example": ["A","Design Phase","","10","15","22"]}), hide_index=True, use_container_width=True)
        st.markdown('<div class="info-box" style="margin-top:1rem;"><strong>Example row:</strong><br><code>B, Build Prototype, A, 3, 6, 9</code><br>Leave predecessors blank for the first activity.</div>', unsafe_allow_html=True)
        st.markdown('<div class="card" style="margin-top:1rem;"><div class="section-label">Interpreting Results</div><p><strong>Mean Duration</strong> — average completion time across all simulations.</p><p><strong>Service Level</strong> — probability of finishing within that duration.</p><p><strong>Critical Path</strong> — longest activity chain; delays here delay everything.</p><p><strong>Std Deviation</strong> — higher means more schedule uncertainty.</p></div>', unsafe_allow_html=True)

elif page == "About PERT":
    st.markdown('<div class="section-label">About PERT</div>', unsafe_allow_html=True)
    col1, col2 = st.columns(2, gap="large")
    with col1:
        st.markdown('''<div class="card card-accent"><div class="section-label">What is PERT?</div>
          <p>PERT (Program Evaluation and Review Technique) is a statistical project management tool developed by the U.S. Navy in 1958. It models uncertainty in task durations using three estimates per activity.</p>
        </div>
        <div class="card"><div class="section-label">Three-Point Estimate</div>
          <p><strong style="color:#0a7e8c;">Minimum (a)</strong> — best-case scenario.</p>
          <p><strong style="color:#0a7e8c;">Average / Most Likely (m)</strong> — realistic estimate.</p>
          <p><strong style="color:#0a7e8c;">Maximum (b)</strong> — worst-case scenario.</p>
          <div class="info-box" style="margin-top:1rem;"><strong>PERT Mean:</strong> (a + 4m + b) / 6<br><strong>Std Dev:</strong> (b - a) / 6</div>
        </div>''', unsafe_allow_html=True)
    with col2:
        st.markdown('''<div class="card card-accent"><div class="section-label">Monte Carlo Simulation</div>
          <p>Instead of computing PERT once, this app runs thousands of simulations, randomly sampling each activity duration from a triangular distribution. The result is a full probability distribution of project durations.</p>
        </div>
        <div class="card"><div class="section-label">Critical Path Method</div>
          <p>The critical path is the longest sequence of dependent activities. Delays on this path directly delay the project end date.</p>
          <div class="info-box" style="margin-top:1rem;"><strong>Service Level:</strong> A 95% service level means there is a 95% probability the project finishes within the stated duration.</div>
        </div>''', unsafe_allow_html=True)

else:
    acts = st.session_state.activities
    unit = st.session_state.time_unit
    left_col, right_col = st.columns([1.05, 1.95], gap="large")

    with left_col:
        itabs = st.tabs(["Activities", "Upload File"])
        with itabs[0]:
            st.markdown('<div class="section-label" style="margin-top:0.5rem;">Current Activities</div>', unsafe_allow_html=True)
            rows = [{"Label": l, "Activity": i["name"], "Pred.": ", ".join(i["predecessors"]) or "-", "Min": i["min"], "Avg": i["avg"], "Max": i["max"]} for l, i in acts.items()]
            st.dataframe(pd.DataFrame(rows), use_container_width=True, hide_index=True, height=220)
            with st.expander("Add / Edit Activity"):
                existing_labels = list(acts.keys())
                new_label = st.text_input("Label", max_chars=5, placeholder="e.g. H").strip().upper()
                new_name  = st.text_input("Activity Name", placeholder="e.g. User Testing")
                selected_preds = st.multiselect("Immediate Predecessors", options=[l for l in existing_labels if l != new_label])
                c1, c2, c3 = st.columns(3)
                with c1: new_min = st.number_input("Min", min_value=0, value=1, key="add_min")
                with c2: new_avg = st.number_input("Avg", min_value=0, value=3, key="add_avg")
                with c3: new_max = st.number_input("Max", min_value=0, value=5, key="add_max")
                if st.button("Save Activity", use_container_width=True):
                    if not new_label: st.error("Enter a label.")
                    elif new_min > new_avg or new_avg > new_max: st.error("Need: Min <= Avg <= Max")
                    elif not new_name: st.error("Enter an activity name.")
                    else:
                        st.session_state.activities[new_label] = {"name": new_name, "predecessors": selected_preds, "min": new_min, "avg": new_avg, "max": new_max}
                        st.session_state.results = None; st.rerun()
            with st.expander("Remove Activity"):
                to_remove = st.selectbox("Select", options=["-"] + list(acts.keys()))
                if st.button("Remove", use_container_width=True) and to_remove != "-":
                    del st.session_state.activities[to_remove]
                    for a in st.session_state.activities:
                        st.session_state.activities[a]["predecessors"] = [p for p in st.session_state.activities[a]["predecessors"] if p != to_remove]
                    st.session_state.results = None; st.rerun()
            if st.button("Reset to Example", use_container_width=True):
                del st.session_state["activities"]; st.session_state.results = None; st.rerun()

        with itabs[1]:
            st.markdown('<div class="info-box">Upload a <strong>CSV or Excel</strong> file with columns: label, name, predecessors, min, avg, max</div>', unsafe_allow_html=True)
            uploaded = st.file_uploader("Upload", type=["csv","xlsx","xls"], label_visibility="collapsed")
            if uploaded:
                try:
                    df_up = pd.read_csv(uploaded) if uploaded.name.endswith(".csv") else pd.read_excel(uploaded)
                    df_up.columns = [c.strip().lower() for c in df_up.columns]
                    if not {"label","name","min","avg","max"}.issubset(set(df_up.columns)):
                        st.error("Missing required columns!")
                    else:
                        new_acts = {}
                        for _, row in df_up.iterrows():
                            lbl = str(row["label"]).strip().upper()
                            preds_raw = str(row.get("predecessors","")).strip()
                            preds = [p.strip().upper() for p in preds_raw.split(",") if p.strip()] if preds_raw and preds_raw.lower() != "nan" else []
                            new_acts[lbl] = {"name": str(row["name"]).strip(), "predecessors": preds, "min": float(row["min"]), "avg": float(row["avg"]), "max": float(row["max"])}
                        st.session_state.activities = new_acts; st.session_state.results = None
                        st.success(f"Loaded {len(new_acts)} activities!"); st.rerun()
                except Exception as e: st.error(f"Error: {e}")
            st.markdown("---")
            template_df = pd.DataFrame([{"label":"A","name":"Design","predecessors":"","min":16,"avg":21,"max":26},{"label":"B","name":"Build Prototype","predecessors":"A","min":3,"avg":6,"max":9}])
            st.download_button("Download CSV Template", data=template_df.to_csv(index=False).encode(), file_name="pert_template.csv", mime="text/csv", use_container_width=True)

    with right_col:
        if run_btn:
            with st.spinner("Running Monte Carlo simulation..."):
                durations = simulate_project(st.session_state.activities, n_simulations=n_sim)
                pert_dur, crit_path, exp_durs, ef = compute_pert_critical_path(st.session_state.activities)
                st.session_state.results = {"durations": durations, "pert_dur": pert_dur, "crit_path": crit_path, "service_level": service_level}

        if st.session_state.results:
            res = st.session_state.results; durations = res["durations"]; sl = res["service_level"]
            mean_dur = np.mean(durations); std_dur = np.std(durations); sl_val = np.percentile(durations, sl)
            st.markdown('<div class="section-label">Forecast Results</div>', unsafe_allow_html=True)
            st.markdown(f'''<div class="metric-grid">
              <div class="metric-tile"><div class="metric-val">{mean_dur:.1f}<span class="metric-unit"> {unit[:2].lower()}</span></div><div class="metric-lbl">Mean Duration</div></div>
              <div class="metric-tile"><div class="metric-val">{std_dur:.1f}<span class="metric-unit"> {unit[:2].lower()}</span></div><div class="metric-lbl">Std Deviation</div></div>
              <div class="metric-tile"><div class="metric-val">{sl_val:.1f}<span class="metric-unit"> {unit[:2].lower()}</span></div><div class="metric-lbl">{sl}% Service Level</div></div>
              <div class="metric-tile"><div class="metric-val">{res["pert_dur"]:.1f}<span class="metric-unit"> {unit[:2].lower()}</span></div><div class="metric-lbl">PERT Expected</div></div>
            </div>''', unsafe_allow_html=True)
            acts_ref = st.session_state.activities
            path_str = " > ".join(f"{l} ({acts_ref[l]['name']})" for l in res["crit_path"])
            st.markdown(f'''<div class="critical-path-box"><div class="critical-path-label">Critical Path</div>
              <div class="critical-path-text">{path_str}</div></div>''', unsafe_allow_html=True)
            rtabs = st.tabs(["Histogram", "Network Diagram", "Activity Breakdown", "Interpretation"])
            with rtabs[0]:
                fig, ax = plt.subplots(figsize=(9, 4))
                fig.patch.set_facecolor("#ffffff"); ax.set_facecolor("#f4f9fb")
                n_bins = min(60, int(np.sqrt(len(durations))))
                counts, bin_edges, patches = ax.hist(durations, bins=n_bins, color=TEAL, edgecolor="#ffffff", linewidth=0.5, alpha=0.88)
                for patch, left in zip(patches, bin_edges[:-1]):
                    if left >= sl_val: patch.set_facecolor(RED); patch.set_alpha(0.75)
                ax.axvline(mean_dur, color="#0a7e8c", linewidth=2.2, linestyle="--", label=f"Mean: {mean_dur:.1f}")
                ax.axvline(sl_val, color=RED, linewidth=2.2, linestyle=":", label=f"{sl}th pct: {sl_val:.1f}")
                ax.set_xlabel(f"Project Duration ({unit})", color="#4a6a72", fontsize=12)
                ax.set_ylabel("Frequency", color="#4a6a72", fontsize=12)
                ax.tick_params(colors="#4a6a72", labelsize=11)
                for spine in ax.spines.values(): spine.set_edgecolor("#e0eef1")
                ax.legend(facecolor="#ffffff", edgecolor="#e0eef1", labelcolor="#1a2332", fontsize=11)
                ax.set_title("Distribution of Simulated Project Durations", fontsize=13, color="#1a2332", fontweight="bold", pad=10)
                fig.tight_layout(pad=1.5); st.pyplot(fig); plt.close(fig)
            with rtabs[1]:
                net_fig = draw_network_diagram(st.session_state.activities, res["crit_path"])
                st.pyplot(net_fig); plt.close(net_fig)
            with rtabs[2]:
                breakdown = [{"Label": lbl, "Activity": info["name"], "Predecessors": ", ".join(info["predecessors"]) or "-", f"Min ({unit})": info["min"], f"Avg ({unit})": info["avg"], f"Max ({unit})": info["max"], "PERT Mean": round(pert_mean(info["min"],info["avg"],info["max"]),2), "Std Dev": round(pert_std(info["min"],info["max"]),2), "Critical": "YES" if lbl in res["crit_path"] else ""} for lbl, info in acts_ref.items()]
                st.dataframe(pd.DataFrame(breakdown), use_container_width=True, hide_index=True)
            with rtabs[3]:
                st.markdown(f'''<div class="card card-accent">
                  <div class="section-label">Summary and Recommendations</div>
                  <p>Ran <strong>{n_sim:,} simulations</strong> in <strong>{unit.lower()}</strong>.</p>
                  <p>Expected duration: <strong style="color:{TEAL}">{mean_dur:.1f} {unit.lower()}</strong>. For <strong style="color:{RED}">{sl}% on-time probability</strong>, plan for <strong style="color:{RED}">{sl_val:.1f} {unit.lower()}</strong>.</p>
                  <p>Critical path: <strong style="color:{RED}">{" > ".join(res["crit_path"])}</strong>. Focus risk management on these activities.</p>
                  <p>Non-critical activities have schedule float and can have resources reallocated without affecting the end date.</p>
                </div>''', unsafe_allow_html=True)
            st.markdown("---")
            st.markdown('<div class="section-label">Export Results</div>', unsafe_allow_html=True)
            exp_col1, exp_col2 = st.columns(2)
            with exp_col1:
                try:
                    excel_data = export_excel(acts_ref, durations, res["pert_dur"], res["crit_path"], sl, unit)
                    st.download_button("Download Excel Report", data=excel_data, file_name="projectpert_results.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", use_container_width=True)
                except Exception as e: st.error(f"Excel error: {e}")
            with exp_col2:
                try:
                    word_data = export_word(acts_ref, durations, res["pert_dur"], res["crit_path"], sl, unit)
                    st.download_button("Download Word Report", data=word_data, file_name="projectpert_report.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
                except Exception as e: st.error(f"Word error: {e}")
        else:
            st.markdown(f'''<div style="height:480px;display:flex;flex-direction:column;align-items:center;justify-content:center;border:2px dashed #b2dde3;border-radius:20px;background:#ffffff;text-align:center;gap:1.2rem;">
              <div style="font-size:4rem;">◈</div>
              <div style="font-family:serif;font-size:1.6rem;color:#0a7e8c;font-weight:600;">Ready to Forecast</div>
              <div style="font-size:0.95rem;color:#6b8a92;max-width:340px;line-height:1.7;">Add activities on the left, configure settings in the sidebar, then click <strong style="color:{TEAL}">Run Simulation</strong>.</div>
            </div>''', unsafe_allow_html=True)
